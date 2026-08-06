import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Set Margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("MediAssist AI — Technical Documentation")
run.font.name = "Calibri"
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = subtitle.add_run("Intelligent Multi-Agent Healthcare Assistant Framework")
run_sub.font.name = "Calibri"
run_sub.font.size = Pt(13)
run_sub.font.italic = True
run_sub.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph() # Spacer

# Heading 1: Executive Summary
h1 = doc.add_heading("1. Executive Summary", level=1)
p = doc.add_paragraph(
    "MediAssist AI is an enterprise-grade multi-agent healthcare assistant constructed "
    "using LangGraph, LangChain, and Streamlit. The system automates routine clinical and "
    "administrative patient interactions—including appointment management, symptom triage, "
    "medication inquiries, and laboratory report analysis—while enforcing strict medical safety "
    "guardrails through a dedicated reflection layer."
)

# Heading 2: System Architecture & Workflow
doc.add_heading("2. System Architecture & Workflow", level=1)
doc.add_heading("2.1 Multi-Agent State Graph Workflow", level=2)
doc.add_paragraph(
    "The system uses a Directed Acyclic Graph (DAG) state workflow built on LangGraph. "
    "All user interactions pass through a central Supervisor node that classifies intent dynamically "
    "and routes the conversation to specialized sub-agents. Responses from domain agents flow "
    "through a mandatory Reflection Node before being returned to the user interface."
)

doc.add_heading("2.2 Node Responsibilities", level=2)
nodes = [
    ("Supervisor Agent (supervisor.py)", "Uses hybrid keyword matching and LLM intent classification to dynamically route patient queries turn-by-turn."),
    ("Appointment Agent (agents/appointment_agent.py)", "Handles appointment scheduling, rescheduling, and cancellation with Human-in-the-Loop (HITL) checks."),
    ("View Appointment Agent (agents/view_appointment_agent.py)", "Fetches and displays existing patient bookings without triggering new booking workflows."),
    ("Symptom Agent (agents/symptom_agent.py)", "Identifies emergency red-flag symptoms (e.g., severe chest pain) and provides non-diagnostic homecare suggestions."),
    ("Medication Agent (agents/medication_agent.py)", "Explains drug usage, dosages, precautions, and cross-checks medications against patient EHR allergies."),
    ("Report Agent (agents/report_agent.py)", "Analyzes structured JSON lab data and unstructured PDF lab reports to highlight out-of-range parameters."),
    ("Reflection Node (agents/reflection_node.py)", "Filters output artifacts, strips unverified diagnostic claims, and appends mandatory medical disclaimers.")
]

for title_str, desc in nodes:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f"{title_str}: ")
    r.bold = True
    p.add_run(desc)

# Heading 3: Key Architectural Features
doc.add_heading("3. Key Architectural Features", level=1)

doc.add_heading("3.1 Shared Graph State (memory/state.py)", level=2)
doc.add_paragraph("State transitions across nodes use a centralized AgentState schema containing:")
state_items = [
    ("messages", "Appended message history for multi-turn conversations."),
    ("patient_id", "Tracks the active EHR patient context."),
    ("current_intent & next_node", "Set dynamically by the Supervisor node."),
    ("agent_outputs", "Staging store for individual sub-agent outputs."),
    ("pdf_context", "Retrieved context from uploaded medical PDF reports."),
    ("final_response", "Post-reflection sanitized Markdown output presented to the user.")
]
for item, desc in state_items:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f"{item}: ")
    r.bold = True
    p.add_run(desc)

doc.add_heading("3.2 Document RAG Pipeline (utils/rag_pipeline.py)", level=2)
rag_steps = [
    "Text Extraction: Parses text streams from uploaded PDF lab reports.",
    "Text Chunking: Splits content into 400-character windows (overlap: 50) via RecursiveCharacterTextSplitter.",
    "Embeddings & Vector Store: Dense embeddings generated via sentence-transformers/all-MiniLM-L6-v2 and indexed in a FAISS vector store."
]
for i, step in enumerate(rag_steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_heading("3.3 API Key Load Balancing (utils/llm_factory.py)", level=2)
p = doc.add_paragraph("To optimize performance and eliminate rate limits, requests are distributed across dual API key allocations:")
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run("Groq Key Group 1: ").bold = True
p1.add_run("Primary routing & administrative tasks (supervisor, appointment, medication).")
p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("Groq Key Group 2: ").bold = True
p2.add_run("Secondary & safety validation tasks (symptom, report, reflection).")

# Heading 4: Deployment & Setup
doc.add_heading("4. Deployment & Setup Commands", level=1)

p_env = doc.add_paragraph()
p_env.add_run("Environment Setup (.env):\n").bold = True
p_env.add_run("GROQ_API_KEY_1=your_primary_groq_key\nGROQ_API_KEY_2=your_secondary_groq_key\nVOICE=your_groq_whisper_voice_key")

p_run = doc.add_paragraph()
p_run.add_run("Run Commands:\n").bold = True
p_run.add_run("pip install -r requirements.txt\nstreamlit run app.py")

# Save file
doc.save("MediAssist_AI_Documentation.docx")
print("Successfully generated 'MediAssist_AI_Documentation.docx'!")